from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from pydantic import BaseModel
from typing import List, Optional
from sqlalchemy.ext.asyncio import AsyncSession
from app.db.session import get_db
from app.api.v1.deps import get_current_user, check_analysis_quota, check_drafting_quota
from app.models.user import User
from app.models.subscription import Subscription
from app.services.rag_service import rag_service
import os, tempfile, logging

logger = logging.getLogger(__name__)

MAX_EXTRACT_SIZE = 50 * 1024 * 1024

router = APIRouter()

class AnalysisRequest(BaseModel):
    text: str
    analysis_type: str = "general"

class QueryRequest(BaseModel):
    question: str
    top_k: int = 5

class DraftRequest(BaseModel):
    doc_type: str
    context: str
    instructions: str = ""

class QueryResponse(BaseModel):
    answer: str
    sources: List[dict]
    chunks_used: int

@router.post("/analyze")
async def analyze_legal_text(data: AnalysisRequest, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user), sub: Subscription = Depends(check_analysis_quota)):
    try:
        result = await rag_service.analyze_case(data.text)
        sub.analyses_used += 1
        await db.commit()
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Analysis failed", exc_info=e)
        raise HTTPException(status_code=500, detail="خطأ داخلي في الخادم")

@router.post("/query", response_model=QueryResponse)
async def query_legal(data: QueryRequest, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user), sub: Subscription = Depends(check_analysis_quota)):
    try:
        result = await rag_service.query(data.question, top_k=data.top_k)
        sub.analyses_used += 1
        await db.commit()
        return QueryResponse(**result)
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Query failed", exc_info=e)
        raise HTTPException(status_code=500, detail="خطأ داخلي في الخادم")

@router.post("/draft")
async def draft_document(data: DraftRequest, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user), sub: Subscription = Depends(check_drafting_quota)):
    try:
        result = await rag_service.draft_legal_document(data.doc_type, data.context, data.instructions)
        sub.drafts_used += 1
        await db.commit()
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Drafting failed", exc_info=e)
        raise HTTPException(status_code=500, detail="خطأ داخلي في الخادم")

from fastapi.responses import Response

class DraftToPdfRequest(BaseModel):
    text: str
    title: str = "مستند قانوني"
    doc_type: str = "document"

@router.post("/draft-to-pdf")
async def draft_to_pdf(data: DraftToPdfRequest):
    """Generate a professionally formatted PDF document with Arabic RTL support."""
    try:
        from weasyprint import HTML
        import html as html_mod

        type_labels = {"claim": "صحيفة دعوى", "defense": "مذكرة دفاع", "memo": "مذكرة قانونية", "contract": "عقد", "appeal": "استئناف / تمييز"}
        label = type_labels.get(data.doc_type, "مستند قانوني")

        lines_html = ""
        for line in data.text.split("\n"):
            stripped = line.strip()
            if not stripped:
                lines_html += "<br/>"
            elif stripped.startswith("بسم الله") or stripped.startswith("المادة") or stripped.startswith("الباب") or stripped.startswith("الفصل"):
                lines_html += f"<h3>{html_mod.escape(stripped)}</h3>"
            elif stripped.startswith("═") or stripped.startswith("──") or stripped.startswith("==="):
                lines_html += "<hr/>"
            else:
                lines_html += f"<p>{html_mod.escape(stripped)}</p>"

        font_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "..", "..", "assets")
        font_regular = os.path.join(font_dir, "Tajawal-Regular.ttf")
        font_bold = os.path.join(font_dir, "Tajawal-Bold.ttf")

        font_face = ""
        body_font = "sans-serif"
        title_font = "sans-serif"
        if os.path.exists(font_regular) and os.path.exists(font_bold):
            font_face = f"""
                @font-face {{ font-family: 'Tajawal'; src: url('file:///{font_regular}'); }}
                @font-face {{ font-family: 'Tajawal-Bold'; src: url('file:///{font_bold}'); }}
            """
            body_font = "Tajawal, sans-serif"
            title_font = "Tajawal-Bold, sans-serif"

        today = __import__('datetime').date.today().strftime('%Y/%m/%d')
        doc_num = f"QN-{__import__('random').randint(100000, 999999)}"

        html_str = f"""<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
<meta charset="utf-8">
{font_face}
<style>
  @page {{ size: A4; margin: 2cm 2.5cm; }}
  body {{ font-family: {body_font}; font-size: 12pt; line-height: 1.8; color: #1a1a1a; }}
  h1 {{ font-family: {title_font}; font-size: 18pt; text-align: center; margin-bottom: 0.3cm; color: #1e3a5f; }}
  h3 {{ font-family: {title_font}; font-size: 13pt; color: #1e3a5f; margin: 0.4cm 0 0.2cm 0; }}
  .basmala {{ text-align: center; font-size: 14pt; color: #2d5016; margin-bottom: 0.8cm; }}
  .meta {{ text-align: left; font-size: 10pt; color: #666; margin-bottom: 0.5cm; border-bottom: 1px solid #ddd; padding-bottom: 0.3cm; }}
  .meta span {{ margin-left: 1cm; }}
  p {{ text-align: justify; margin: 0.2cm 0; text-indent: 1cm; }}
  hr {{ border: none; border-top: 2px solid #1e3a5f; margin: 0.6cm 0; }}
  .signatures {{ margin-top: 2cm; display: flex; justify-content: space-between; }}
  .sig-block {{ text-align: center; width: 40%; }}
  .sig-block p {{ font-weight: bold; margin-bottom: 0.2cm; text-indent: 0; }}
  .sig-line {{ border-bottom: 1px solid #333; width: 60%; margin: 0.5cm auto 0.2cm; }}
  .footer {{ text-align: center; font-size: 9pt; color: #999; margin-top: 1cm; border-top: 1px solid #eee; padding-top: 0.3cm; }}
</style>
</head>
<body>

<div class="basmala">بسم الله الرحمن الرحيم</div>

<h1>{html_mod.escape(data.title)}</h1>

<div class="meta">
  <span>النوع: {html_mod.escape(label)}</span>
  <span>رقم المسودة: {doc_num}</span>
  <span>التاريخ: {today}م</span>
  <span>منصة قانوني AI</span>
</div>

<hr/>

{lines_html}

<div class="signatures">
  <div class="sig-block">
    <p>الطرف الأول</p>
    <div class="sig-line"></div>
    <small>التوقيع والختم</small>
  </div>
  <div class="sig-block">
    <p>الطرف الثاني</p>
    <div class="sig-line"></div>
    <small>التوقيع</small>
  </div>
</div>

<div class="footer">تم إنشاء هذا المستند عبر منصة قانوني AI — للمرجعية القانونية وليس بديلاً عن استشارة المحامي المرخص</div>

</body>
</html>"""

        pdf_bytes = HTML(string=html_str).write_pdf()
        return Response(content=pdf_bytes, media_type="application/pdf",
                        headers={"Content-Disposition": f'attachment; filename="{data.title}.pdf"'})
    except Exception as e:
        logger.error("PDF generation failed", exc_info=e)
        raise HTTPException(status_code=500, detail="فشل إنشاء PDF")


class DraftToDocxRequest(BaseModel):
    text: str
    title: str = "مستند قانوني"
    doc_type: str = "document"

@router.post("/draft-to-docx")
async def draft_to_docx(data: DraftToDocxRequest):
    """Generate a professionally formatted DOCX document with Arabic RTL support."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        import io, datetime, random

        type_labels = {"claim": "صحيفة دعوى", "defense": "مذكرة دفاع", "memo": "مذكرة قانونية", "contract": "عقد", "appeal": "استئناف / تمييز"}
        label = type_labels.get(data.doc_type, "مستند قانوني")

        doc = Document()

        # Set RTL and page margins
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # Helper to add RTL paragraph
        def add_rtl_paragraph(text, bold=False, font_size=12, color=None, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(6), first_indent=Cm(1)):
            p = doc.add_paragraph()
            p.alignment = alignment
            p.paragraph_format.space_after = space_after
            p.paragraph_format.first_line_indent = first_indent
            # Set RTL for the run
            from docx.oxml.ns import qn
            pPr = p._p.get_or_add_pPr()
            bidi = pPr.makeelement(qn('w:bidi'), {})
            pPr.append(bidi)
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            run.font.name = 'Traditional Arabic'
            # Set complex script font
            rPr = run._r.get_or_add_rPr()
            cs = rPr.makeelement(qn('w:cs'), {'val': 'Traditional Arabic'})
            rPr.append(cs)
            if bold:
                run.font.bold = True
            if color:
                run.font.color.rgb = RGBColor(*color)
            return p

        # Bismillah
        add_rtl_paragraph("بسم الله الرحمن الرحيم", bold=True, font_size=14,
                         color=(45, 80, 22), alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Cm(0))

        # Title
        add_rtl_paragraph(data.title, bold=True, font_size=18,
                         color=(30, 58, 95), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4), first_indent=Cm(0))

        # Meta info
        today = datetime.date.today().strftime('%Y/%m/%d')
        doc_num = f"QN-{random.randint(100000, 999999)}"
        meta_text = f"النوع: {label}  |  رقم المسودة: {doc_num}  |  التاريخ: {today}م  |  منصة قانوني AI"
        add_rtl_paragraph(meta_text, font_size=10, color=(102, 102, 102),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12), first_indent=Cm(0))

        # Separator
        sep_p = doc.add_paragraph()
        sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml.ns import qn
        pBdr = sep_p._p.get_or_add_pPr()
        bottom = pBdr.makeelement(qn('w:pBdr'), {})
        b = bottom.makeelement(qn('w:bottom'), {'val': 'single', 'sz': '12', 'space': '1', 'color': '1E3A5F'})
        bottom.append(b)
        pBdr.append(bottom)

        # Body text
        for line in data.text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            is_heading = any(stripped.startswith(p) for p in ["بسم الله", "المادة", "الباب", "الفصل", "القسم"])
            is_separator = any(stripped.startswith(s) for s in ["═", "──", "==="])
            if is_separator:
                sep_p2 = doc.add_paragraph()
                sep_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pBdr2 = sep_p2._p.get_or_add_pPr()
                bottom2 = pBdr2.makeelement(qn('w:pBdr'), {})
                b2 = bottom2.makeelement(qn('w:bottom'), {'val': 'single', 'sz': '6', 'space': '1', 'color': 'CCCCCC'})
                bottom2.append(b2)
                pBdr2.append(bottom2)
            elif is_heading:
                add_rtl_paragraph(stripped, bold=True, font_size=13,
                                 color=(30, 58, 95), alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=Cm(0))
            else:
                add_rtl_paragraph(stripped, font_size=12)

        # Signatures
        doc.add_paragraph()  # spacing
        sig_p = doc.add_paragraph()
        sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Footer note
        doc.add_paragraph()
        add_rtl_paragraph("تم إنشاء هذا المستند عبر منصة قانوني AI — للمرجعية القانونية وليس بديلاً عن استشارة المحامي المرخص",
                         font_size=9, color=(153, 153, 153), alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Cm(0))

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return Response(content=buffer.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        headers={"Content-Disposition": f'attachment; filename="{data.title}.docx"'})
    except Exception as e:
        logger.error("DOCX generation failed", exc_info=e)
        raise HTTPException(status_code=500, detail="فشل إنشاء ملف Word")

@router.post("/extract-text")
async def extract_text_from_file(file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename or "file")[1].lower()
    content_bytes = await file.read()
    if len(content_bytes) > MAX_EXTRACT_SIZE:
        raise HTTPException(status_code=413, detail="حجم الملف يتجاوز الحد المسموح (50MB)")
    text = ""

    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content_bytes))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            method = "pdf"
        elif ext == ".docx":
            from docx import Document
            import io
            doc = Document(io.BytesIO(content_bytes))
            text = "\n".join(p.text for p in doc.paragraphs)
            method = "docx"
        elif ext == ".txt":
            text = content_bytes.decode("utf-8", errors="replace")
            method = "txt"
        else:
            raise HTTPException(status_code=400, detail=f"نوع الملف {ext} غير مدعوم. استخدم PDF, DOCX, أو TXT")
    except Exception as e:
        logger.error("Text extraction failed", exc_info=e)
        raise HTTPException(status_code=500, detail="فشل استخراج النص")

    if not text.strip():
        raise HTTPException(status_code=400, detail="لم يتم استخراج أي نص من الملف. قد يكون الملف فارغاً أو محمياً بكلمة مرور")

    return {"text": text, "method": method, "filename": os.path.basename(file.filename or "file")}

@router.post("/ingest")
async def ingest_document(doc_id: str, title: str, content: str, metadata: dict = None, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)):
    if current_user.role.value not in ["admin", "reviewer"]:
        raise HTTPException(status_code=403, detail="صلاحية غير كافية")
    try:
        chunks_count = await rag_service.ingest_document(doc_id, title, content, metadata)
        return {"status": "success", "chunks_indexed": chunks_count}
    except Exception as e:
        logger.error("Document ingestion failed", exc_info=e)
        raise HTTPException(status_code=500, detail="فشل إدراج المستند")
